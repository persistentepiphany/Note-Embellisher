import os
from pathlib import Path
from typing import Optional

try:
    from docx import Document  # type: ignore
    DOCX_AVAILABLE = True
    _DOCX_IMPORT_ERROR = None
except ImportError as e:  # pragma: no cover - optional dependency
    Document = None  # type: ignore
    DOCX_AVAILABLE = False
    _DOCX_IMPORT_ERROR = e


class DocumentExportService:
    """Create DOCX and TXT exports for processed notes."""

    def __init__(self, output_dir: Optional[str] = None) -> None:
        base_dir = output_dir or "generated_pdfs"
        self.output_dir = Path(base_dir)
        self.output_dir.mkdir(exist_ok=True)
        if not DOCX_AVAILABLE:
            print(
                "python-docx not installed. DOCX exports will be disabled until the dependency is added."
            )
            if _DOCX_IMPORT_ERROR:
                print(f"python-docx import error: {_DOCX_IMPORT_ERROR}")

    # ------------------------------------------------------------------
    def generate_docx(self, content: str, title: str, filename_slug: str) -> str:
        if not DOCX_AVAILABLE or not Document:
            raise RuntimeError("DOCX export is not available (missing python-docx)")

        doc = Document()
        doc.add_heading(title, level=0)
        for block in split_paragraphs(content):
            doc.add_paragraph(block)
        output_path = self.output_dir / f"{filename_slug}.docx"
        doc.save(output_path)
        return str(output_path)

    def generate_txt(self, content: str, filename_slug: str) -> str:
        output_path = self.output_dir / f"{filename_slug}.txt"
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(content)
        return str(output_path)


def split_paragraphs(content: str) -> list[str]:
    """Simple helper to keep paragraph spacing in exports."""
    if not content:
        return [""]
    parts = [part.strip() for part in content.split("\n\n")]
    return [part for part in parts if part]


export_service = DocumentExportService()
